from flask import Blueprint, request, jsonify, render_template, session, redirect, url_for, send_file
from config import get_db
from datetime import datetime
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import io
import os
import csv
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import xlsxwriter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4, landscape


preferences_bp = Blueprint("preferences_bp", __name__)

# -------------------------
# API - 志願填寫
# -------------------------
@preferences_bp.route('/fill_preferences', methods=['GET', 'POST'])
def fill_preferences():
    # 1. 登入檢查
    if 'user_id' not in session or session.get('role') != 'student':
        return redirect(url_for('auth_bp.login_page'))

    student_id = session['user_id']

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    message = None

    if request.method == 'POST':
        preferences = []
        for i in range(1, 6):
            company_id = request.form.get(f'company_{i}') 
            if company_id:
              job_id = request.form.get(f'job_{i}')
              preferences.append((student_id, i, company_id, job_id, datetime.now()))

        try:
            # 刪除舊志願
            cursor.execute("DELETE FROM student_preferences WHERE student_id = %s", (student_id,))
            conn.commit()

            # 新增志願
            if preferences:
                cursor.executemany("""
                    INSERT INTO student_preferences (student_id, preference_order, company_id, job_id, submitted_at)
                    VALUES (%s, %s, %s, %s, %s)
                """, preferences)
                conn.commit()
                message = "✅ 志願序已成功送出"
            else:
                message = "⚠️ 未選擇任何志願，公司清單已重置"
        except Exception as e:
            print("寫入志願錯誤：", e)
            message = "❌ 發生錯誤，請稍後再試"

    # 不管是 GET 還是 POST，都要載入公司列表及該學生已填的志願
    cursor.execute("SELECT id, company_name FROM internship_companies WHERE status = 'approved'")
    companies = cursor.fetchall()

    cursor.execute("""
        SELECT preference_order, company_id 
        FROM student_preferences 
        WHERE student_id = %s 
        ORDER BY preference_order
    """, (student_id,))
    prefs = cursor.fetchall()

    cursor.close()
    conn.close()

    # 把 prefs 轉成 list，index 對應志願順序 -1
    submitted_preferences = [None] * 5
    for pref in prefs:
        order = pref['preference_order']
        company_id = pref['company_id']
        if 1 <= order <= 5:
            submitted_preferences[order - 1] = company_id

    return render_template('preferences/fill_preferences.html',
        companies=companies,
        submitted_preferences=submitted_preferences,
        message=message
    )

# -------------------------
# API - 選擇角色
# -------------------------
@preferences_bp.route('/api/select_role', methods=['POST'])
def select_role():
    data = request.json
    username = data.get("username")
    role = data.get("role")

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id FROM users WHERE username=%s AND role=%s", (username, role))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

    if user:
        session["user_id"] = user["id"]
        session["role"] = role
        return jsonify({"success": True})
    else:
        return jsonify({"success": False, "message": "無此角色"}), 404


# -------------------------
# 班導查看志願序
# -------------------------
@preferences_bp.route('/review_preferences')
def review_preferences():
    if 'user_id' not in session or session.get('role') not in ['teacher', 'director']:
      return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導
        cursor.execute("""
            SELECT c.id AS class_id
            FROM classes c
            JOIN classes_teacher ct ON c.id = ct.class_id
            WHERE ct.teacher_id = %s AND ct.role = '班導師'
        """, (user_id,))
        class_info = cursor.fetchone()
        if not class_info:
            return "你不是班導，無法查看志願序", 403

        class_id = class_info['class_id']

        # 查詢班上學生及其志願
        cursor.execute("""
            SELECT 
                u.id AS student_id,
                u.name AS student_name,
                sp.preference_order,
                ic.company_name,
                ij.title AS job_title,
                sp.submitted_at        
            FROM users u
            LEFT JOIN student_preferences sp ON u.id = sp.student_id
            LEFT JOIN internship_companies ic ON sp.company_id = ic.id
            LEFT JOIN internship_jobs ij ON sp.job_id = ij.id       
            WHERE u.class_id = %s AND u.role = 'student'
            ORDER BY u.name, sp.preference_order
        """, (class_id,))
        results = cursor.fetchall()

        # 整理資料結構給前端使用
        student_data = defaultdict(list)
        for row in results:
            if row['preference_order'] and row['company_name']:
                student_data[row['student_name']].append({
                    'order': row['preference_order'],
                    'company': row['company_name'],
                    'job_title': row['job_title'],
                    'submitted_at': row['submitted_at']
                })

        return render_template('preferences/review_preferences.html', student_data=student_data)

    except Exception as e:
        print("取得志願資料錯誤：", e)
        return "伺服器錯誤", 500
    finally:
        cursor.close()
        conn.close()


# -------------------------
# Excel 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_excel')
def export_preferences_excel():
    if 'user_id' not in session or session.get('role') not in ['teacher', 'director']:
       return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導
        cursor.execute("""
        SELECT c.id AS class_id, c.name AS class_name
        FROM classes c
        JOIN classes_teacher ct ON c.id = ct.class_id
        WHERE ct.teacher_id = %s AND ct.role = '班導師'
        """, (user_id,))
        class_info = cursor.fetchone()
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 查詢班上學生及其志願
        cursor.execute("""
            SELECT 
                u.id AS student_id,
                u.name AS student_name,
                u.username AS student_number, 
                sp.preference_order,
                ic.company_name,
                sp.submitted_at
            FROM users u
            LEFT JOIN student_preferences sp ON u.id = sp.student_id
            LEFT JOIN internship_companies ic ON sp.company_id = ic.id
            WHERE u.class_id = %s AND u.role = 'student'
            ORDER BY u.name, sp.preference_order
        """, (class_id,))
        results = cursor.fetchall()

        # 創建 Excel 工作簿
        wb = Workbook()
        ws = wb.active
        ws.title = f"{class_name}志願序"

        # 設定樣式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="0066CC", end_color="0066CC", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 寫入標題
        ws.merge_cells('A1:G1')
        title_cell = ws['A1']
        title_cell.value = f"{class_name} - 學生實習志願序統計表"
        title_cell.font = Font(bold=True, size=16, color="0066CC")
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        # 寫入日期
        ws.merge_cells('A2:G2')
        date_cell = ws['A2']
        date_cell.value = f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        date_cell.alignment = Alignment(horizontal="center")

        # 設定表頭
        headers = ['學生姓名', '學號', '第一志願', '第二志願', '第三志願', '第四志願', '第五志願']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border

        # 整理學生資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'preferences': [''] * 5,
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row['student_name']
            if student_name:
                student_data[student_name]['name'] = student_name
                student_data[student_name]['student_number'] = row['student_number'] or ''
                
                if row['preference_order'] and row['company_name']:
                    order = row['preference_order'] - 1  # 轉為 0-based index
                    if 0 <= order < 5:
                        student_data[student_name]['preferences'][order] = row['company_name']
                        if row['submitted_at']:
                            student_data[student_name]['submitted_times'][order] = row['submitted_at'].strftime('%m/%d %H:%M')

        # 寫入學生資料
        row_num = 5
        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            
            # 學生姓名
            ws.cell(row=row_num, column=1, value=data['name']).border = border
            # 學號
            ws.cell(row=row_num, column=2, value=data['student_number']).border = border
            
            # 志願序
            for i in range(5):
                pref_text = data['preferences'][i]
                if pref_text and data['submitted_times'][i]:
                    pref_text += f"\n({data['submitted_times'][i]})"
                
                cell = ws.cell(row=row_num, column=3+i, value=pref_text)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
            
            row_num += 1

        # 添加統計資訊
        ws.cell(row=row_num + 1, column=1, value="統計資訊：").font = Font(bold=True)
        
        # 統計各公司被選擇次數
        company_counts = defaultdict(int)
        for data in student_data.values():
            for pref in data['preferences']:
                if pref:
                    company_counts[pref] += 1

        stats_row = row_num + 2
        ws.cell(row=stats_row, column=1, value="公司名稱").font = Font(bold=True)
        ws.cell(row=stats_row, column=2, value="被選擇次數").font = Font(bold=True)
        
        stats_row += 1
        for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
            ws.cell(row=stats_row, column=1, value=company)
            ws.cell(row=stats_row, column=2, value=count)
            stats_row += 1

        # 調整欄寬
        column_widths = [15, 12, 20, 20, 20, 20, 20]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # 設定行高
        for row in range(5, row_num):
            ws.row_dimensions[row].height = 40

        # 保存到記憶體
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

        # 生成檔案名稱
        filename = f"{class_name}_學生志願序_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        return send_file(
            excel_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        print("導出 Excel 錯誤：", e)
        return "伺服器錯誤", 500
    finally:
        cursor.close()
        conn.close()


# -------------------------
# word 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_word')
def export_preferences_word():
    if 'user_id' not in session or session.get('role') not in ['teacher', 'director']:
       return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導
        cursor.execute("""
          SELECT c.id AS class_id, c.name AS class_name
          FROM classes c
          JOIN classes_teacher ct ON c.id = ct.class_id
          WHERE ct.teacher_id = %s AND ct.role = '班導師'
        """, (user_id,))
        class_info = cursor.fetchone()
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 查詢學生志願
        cursor.execute("""
            SELECT 
                u.id AS student_id,
                u.name AS student_name,
                u.username AS student_number, 
                sp.preference_order,
                ic.company_name,
                sp.submitted_at
            FROM users u
            LEFT JOIN student_preferences sp ON u.id = sp.student_id
            LEFT JOIN internship_companies ic ON sp.company_id = ic.id
            WHERE u.class_id = %s AND u.role = 'student'
            ORDER BY u.name, sp.preference_order
        """, (class_id,))
        results = cursor.fetchall()

        # 整理資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'preferences': [''] * 5,
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row['student_name']
            if student_name:
                student_data[student_name]['name'] = student_name
                student_data[student_name]['student_number'] = row['student_number'] or ''
                if row['preference_order'] and row['company_name']:
                    order = row['preference_order'] - 1
                    if 0 <= order < 5:
                        student_data[student_name]['preferences'][order] = row['company_name']
                        if row['submitted_at']:
                            student_data[student_name]['submitted_times'][order] = row['submitted_at'].strftime('%m/%d %H:%M')

        # 建立 Word 文件
        doc = Document()
        title = doc.add_heading(f"{class_name} - 學生實習志願序統計表", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph("")

        # 學生表格
        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ['學生姓名', '學號', '第一志願', '第二志願', '第三志願', '第四志願', '第五志願']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            row = table.add_row().cells
            row[0].text = data['name']
            row[1].text = data['student_number']
            for i in range(5):
                pref_text = data['preferences'][i]
                if pref_text and data['submitted_times'][i]:
                    pref_text += f"\n({data['submitted_times'][i]})"
                row[2+i].text = pref_text

        doc.add_paragraph("")
        doc.add_heading("統計資訊", level=1)

        # 統計資訊
        company_counts = defaultdict(int)
        for data in student_data.values():
            for pref in data['preferences']:
                if pref:
                    company_counts[pref] += 1

        if company_counts:
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = "Table Grid"
            stats_table.rows[0].cells[0].text = "公司名稱"
            stats_table.rows[0].cells[1].text = "被選擇次數"
            for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
                row = stats_table.add_row().cells
                row[0].text = company
                row[1].text = str(count)

        # 匯出檔案
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"{class_name}_學生志願序_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("Word 匯出錯誤：", e)
        return "伺服器錯誤", 500
    finally:
        cursor.close()
        conn.close()


# -------------------------
# PDF 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_pdf')
def export_preferences_pdf():
    if 'user_id' not in session or session.get('role') not in ['teacher', 'director']:
       return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導
        cursor.execute("""
        SELECT c.id AS class_id, c.name AS class_name
        FROM classes c
        JOIN classes_teacher ct ON c.id = ct.class_id
        WHERE ct.teacher_id = %s AND ct.role = '班導師'
        """, (user_id,))
        class_info = cursor.fetchone()
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 查詢班上學生及其志願，包含公司與職缺與聯絡資訊
        cursor.execute("""
            SELECT 
                u.id AS student_id,
                u.name AS student_name,
                u.username AS student_number,
                u.class_id,
                sp.preference_order,
                sp.submitted_at,
                ic.id AS company_id,
                ic.company_name,
                ic.company_address,
                ic.contact_name,
                ic.contact_phone,
                ic.contact_email,
                ij.id AS job_id,
                ij.title AS job_title
            FROM users u
            LEFT JOIN student_preferences sp ON u.id = sp.student_id
            LEFT JOIN internship_companies ic ON sp.company_id = ic.id
            LEFT JOIN internship_jobs ij ON sp.job_id = ij.id
            WHERE u.class_id = %s AND u.role = 'student'
            ORDER BY u.name, sp.preference_order
        """, (class_id,))
        results = cursor.fetchall()

        # 整理學生資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'class_id': '',
            'preferences': [None] * 5,  # each entry will be dict or None
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row.get('student_name')
            if not student_name:
                continue
            student = student_data[student_name]
            student['name'] = student_name
            student['student_number'] = row.get('student_number') or ''
            student['class_id'] = row.get('class_id') or ''

            pref_order = row.get('preference_order')
            if pref_order and row.get('company_name'):
                idx = pref_order - 1
                if 0 <= idx < 5:
                    student['preferences'][idx] = {
                        'company_name': row.get('company_name') or '',
                        'job_title': row.get('job_title') or row.get('job_title') or '',
                        'company_address': row.get('company_address') or '',
                        'contact_name': row.get('contact_name') or '',
                        'contact_phone': row.get('contact_phone') or '',
                        'contact_email': row.get('contact_email') or ''
                    }
                    if row.get('submitted_at'):
                        student['submitted_times'][idx] = row['submitted_at'].strftime('%Y/%m/%d %H:%M')

        # 準備 PDF（橫式）
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=landscape(A4),
            leftMargin=0.5*inch,
            rightMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=1,  # centered
            textColor=colors.HexColor('#0066CC')
        )
        normal_style = ParagraphStyle(
            'NormalWrap',
            parent=styles['Normal'],
            fontSize=9,
            leading=11
        )

        story = []
        # 標題與日期
        story.append(Paragraph(f"{class_name} - 學生實習志願序統計表", title_style))
        story.append(Paragraph(f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}", normal_style))
        story.append(Spacer(1, 12))

        # 建表格資料（欄位：學生姓名、學號、班級、志願序、公司、職缺、公司地址、聯絡人、聯絡電話、提交時間）
        headers = ['學生姓名', '學號', '班級', '志願序', '公司名稱', '職缺', '公司地址', '聯絡人', '聯絡電話', '提交時間']
        table_data = [headers]

        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            for idx in range(5):
                pref = data['preferences'][idx]
                if not pref:
                    # 如果該志願未填，仍列出空白的該列（可選）
                    row = [
                        data['name'],
                        data['student_number'],
                        class_name,
                        f"第{idx+1}志願",
                        '', '', '', '', '', ''
                    ]
                else:
                    contact = pref.get('contact_name') or ''
                    phone = pref.get('contact_phone') or pref.get('contact_email') or ''
                    row = [
                        data['name'],
                        data['student_number'],
                        class_name,
                        f"第{idx+1}志願",
                        pref.get('company_name', ''),
                        pref.get('job_title', ''),
                        pref.get('company_address', ''),
                        contact,
                        phone,
                        data['submitted_times'][idx] or ''
                    ]
                # 使用 Paragraph 讓長文本可以自動換行
                row = [Paragraph(str(cell), normal_style) for cell in row]
                table_data.append(row)

        # 如果沒有任何學生，顯示提示
        if len(table_data) == 1:
            table_data.append([Paragraph("沒有可顯示的資料", normal_style)] + [''] * (len(headers) - 1))

        # 設定欄寬（橫式需要寬欄）
        col_widths = [1.4*inch, 0.9*inch, 0.9*inch, 0.8*inch, 2.2*inch, 1.6*inch, 2.6*inch, 1.2*inch, 1.2*inch, 1.0*inch]

        table = Table(table_data, colWidths=col_widths, repeatRows=1)

        table_style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 10),
            ('FONTSIZE', (0,1), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F9F9F9')]),
        ])
        table.setStyle(table_style)

        story.append(table)
        story.append(Spacer(1, 12))

        # 統計資訊：計算被選擇次數（以公司+職缺為 key）
        company_counts = defaultdict(int)
        for student in student_data.values():
            for pref in student['preferences']:
                if pref:
                    key = (pref.get('company_name',''), pref.get('job_title',''))
                    company_counts[key] += 1

        if company_counts:
            story.append(Paragraph("統計資訊：公司(職缺) 被選擇次數", styles['Heading3']))
            stats_table_data = [['公司名稱', '職缺', '被選擇次數']]
            for (company, job), count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
                stats_table_data.append([Paragraph(company or '', normal_style),
                                         Paragraph(job or '', normal_style),
                                         Paragraph(str(count), normal_style)])
            stats_table = Table(stats_table_data, colWidths=[3*inch, 2.5*inch, 1*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0066CC')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ]))
            story.append(stats_table)

        # 生成 PDF
        doc.build(story)

        pdf_buffer.seek(0)
        filename = f"{class_name}_學生志願序_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        print("導出 PDF 錯誤：", e)
        return "伺服器錯誤", 500
    finally:
        cursor.close()
        conn.close()
